"""周报导出 - PDF/Word"""
import io, markdown, tempfile
from flask import Blueprint, request, jsonify, send_file

export_bp = Blueprint("export", __name__)


def _md_to_html(md_text):
    body = markdown.markdown(md_text, extensions=["tables"])
    return f"""<html><head><meta charset="utf-8"><style>
body{{font-family:"Noto Sans SC","PingFang SC",sans-serif;max-width:800px;margin:40px auto;line-height:1.8;font-size:14px}}
h1{{color:#1890ff;border-bottom:2px solid #1890ff;padding-bottom:8px}}
h2{{color:#333;margin-top:24px}} ul,ol{{padding-left:20px}}
table{{border-collapse:collapse;width:100%}} th,td{{border:1px solid #ddd;padding:6px 10px;text-align:left}}
</style></head><body>{body}</body></html>"""


@export_bp.route("/api/weekly/export", methods=["POST"])
def weekly_export():
    body = request.json or {}
    content_md = body.get("content", "")
    title = body.get("title", "周报")
    fmt = body.get("format", "pdf")  # pdf or docx

    if not content_md:
        return jsonify({"error": "content 不能为空"}), 400

    if fmt == "pdf":
        from weasyprint import HTML
        html_str = _md_to_html(content_md)
        pdf_bytes = HTML(string=html_str).write_pdf()
        buf = io.BytesIO(pdf_bytes)
        buf.seek(0)
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True, download_name=f"{title}.pdf")

    elif fmt == "docx":
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(title, level=0)
        for line in content_md.split("\n"):
            line = line.rstrip()
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=f"{title}.docx")

    return jsonify({"error": "format 仅支持 pdf/docx"}), 400
